import json
import zipfile
from io import BytesIO

import openpyxl
import pypdf
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


class UnsupportedSubmissionFormat(Exception):
    """The uploaded bytes carry no extractable text."""


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    lines = []
    # Walk the body rather than .paragraphs + .tables, which would append every
    # table after all prose. Review comments anchor to line numbers, so the
    # extracted order has to match what the student sees in Word.
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            lines.append(Paragraph(child, document).text)
        elif child.tag.endswith("}tbl"):
            for row in Table(child, document).rows:
                lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _pdf_text(content: bytes) -> str:
    reader = pypdf.PdfReader(BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if reader.pages and not text.strip():
        # A scanned PDF has no text layer. Returning "" would set is_empty on a
        # real submission, so refuse it instead of grading a blank string.
        raise UnsupportedSubmissionFormat("PDF has no text layer (scan?) — OCR required")
    return text


def _xlsx_text(content: bytes) -> str:
    workbook = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
    try:
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                lines.append("\t".join("" if value is None else str(value) for value in row))
        return "\n".join(lines)
    finally:
        workbook.close()


def _ipynb_text(text: str) -> str | None:
    """Extract cell sources, or None if this text isn't a notebook.

    Detection parses the JSON instead of scanning for keys: "nbformat" is
    written last, so it sits at the very end of a multi-hundred-KB file.
    """
    if text.lstrip()[:1] != "{":
        return None
    try:
        notebook = json.loads(text)
    except (json.JSONDecodeError, RecursionError):
        return None
    if not isinstance(notebook, dict) or "cells" not in notebook or "nbformat" not in notebook:
        return None

    lines = []
    for cell in notebook["cells"]:
        source = cell.get("source", "")
        body = "".join(source) if isinstance(source, list) else source
        lines.append(f"# --- {cell.get('cell_type', 'unknown')} ---")
        lines.append(body.rstrip("\n"))
    return "\n".join(lines)


def _is_binary(content: bytes) -> bool:
    sample = content[:8192]
    if b"\x00" in sample:
        return True
    printable = sum(1 for byte in sample if byte >= 32 or byte in (9, 10, 13))
    return printable / len(sample) < 0.85


def parse_submission_text(content: bytes) -> str:
    """Extract reviewable plain text from an uploaded submission.

    Dispatches on magic bytes: a renamed file or a browser-supplied
    content_type would both lie about the real format.
    """
    if not content:
        return ""

    if content.startswith(b"%PDF"):
        return _pdf_text(content)

    if content.startswith(b"PK\x03\x04"):
        try:
            entries = set(zipfile.ZipFile(BytesIO(content)).namelist())
        except zipfile.BadZipFile as exc:
            raise UnsupportedSubmissionFormat("Corrupt Office/ZIP archive") from exc
        if "word/document.xml" in entries:
            return _docx_text(content)
        if "xl/workbook.xml" in entries:
            return _xlsx_text(content)
        raise UnsupportedSubmissionFormat("ZIP archives other than .docx/.xlsx are not supported")

    if _is_binary(content):
        raise UnsupportedSubmissionFormat("Binary file with no extractable text")

    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        # Coursework predating UTF-8 still shows up as cp1251; _is_binary above
        # keeps this from silently "decoding" arbitrary bytes.
        try:
            text = content.decode("cp1251")
        except UnicodeDecodeError as exc:
            raise UnsupportedSubmissionFormat("Unrecognized text encoding") from exc

    notebook = _ipynb_text(text)
    return text if notebook is None else notebook
